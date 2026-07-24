#!/usr/bin/env python3
"""Validate the generated thesis review DOCX and its rendered PDF."""

from __future__ import annotations

import argparse
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from pypdf import PdfReader


EXPECTED_TITLE = "Reusable Human Judgment in AI-Assisted Domain Model Assessment"
FORBIDDEN_TEXT = {
    "markdown emphasis marker": re.compile(r"\*\*"),
    "markdown fence": re.compile(r"```"),
    "local URI": re.compile(r"file:///|codex-file-citation|oai-mem-citation", re.I),
    "placeholder": re.compile(r"\b(?:TODO|TBD|PLACEHOLDER|LOREM IPSUM)\b", re.I),
}


def validate_docx(path: Path) -> list[str]:
    errors: list[str] = []
    if not path.is_file():
        return [f"DOCX does not exist: {path}"]

    document = Document(path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    normalized_text = re.sub(r"\s+", " ", text)
    if EXPECTED_TITLE not in normalized_text:
        errors.append("expected thesis title is missing")
    for name, pattern in FORBIDDEN_TEXT.items():
        if pattern.search(text):
            errors.append(f"forbidden {name} remains in document text")

    if len(document.inline_shapes) < 4:
        errors.append("fewer than four generated review figures are embedded")
    if len(document.tables) < 10:
        errors.append("expected thesis tables are missing")

    section = document.sections[0]
    expected = {
        "page_width": 8.5,
        "page_height": 11.0,
        "top_margin": 1.0,
        "bottom_margin": 1.0,
        "left_margin": 1.0,
        "right_margin": 1.0,
    }
    for attribute, inches in expected.items():
        actual = getattr(section, attribute).inches
        if abs(actual - inches) > 0.01:
            errors.append(f"{attribute} is {actual:.3f} inches; expected {inches:.3f}")

    for table_index, table in enumerate(document.tables, start=1):
        tbl_pr = table._tbl.tblPr
        layout = tbl_pr.first_child_found_in("w:tblLayout")
        if layout is None or layout.get(qn("w:type")) != "fixed":
            errors.append(f"table {table_index} is not fixed-layout")
        tbl_w = tbl_pr.first_child_found_in("w:tblW")
        if tbl_w is None or tbl_w.get(qn("w:type")) != "dxa":
            errors.append(f"table {table_index} has no explicit DXA width")
        elif int(tbl_w.get(qn("w:w"))) > 9360:
            errors.append(f"table {table_index} exceeds the 9360 DXA content width")
        for row_index, row in enumerate(table.rows, start=1):
            tr_pr = row._tr.get_or_add_trPr()
            if tr_pr.find(qn("w:cantSplit")) is None:
                errors.append(
                    f"table {table_index} row {row_index} may split across pages"
                )
            for height in tr_pr.findall(qn("w:trHeight")):
                if height.get(qn("w:hRule")) == "exact":
                    errors.append(
                        f"table {table_index} row {row_index} uses an exact height"
                    )
        first_tr_pr = table.rows[0]._tr.get_or_add_trPr()
        if first_tr_pr.find(qn("w:tblHeader")) is None:
            errors.append(f"table {table_index} header is not marked to repeat")

    with zipfile.ZipFile(path) as archive:
        document_xml = archive.read("word/document.xml")
        for token, description in (
            (b'w:tblLayout w:type="fixed"', "fixed table layout"),
            (b"<w:cantSplit", "non-splitting table rows"),
            (b"<w:tblHeader", "repeating table headers"),
            (b"<w:keepNext", "keep-with-next pagination"),
        ):
            if token not in document_xml:
                errors.append(f"OOXML is missing {description}")

    return errors


def validate_pdf(path: Path) -> list[str]:
    errors: list[str] = []
    if not path.is_file():
        return [f"PDF does not exist: {path}"]
    pdf = PdfReader(path)
    if not 60 <= len(pdf.pages) <= 120:
        errors.append(f"unexpected PDF page count: {len(pdf.pages)}")
    for page_number, page in enumerate(pdf.pages, start=1):
        text = page.extract_text() or ""
        if len(text.strip()) < 25:
            errors.append(f"PDF page {page_number} appears blank")
        for name, pattern in FORBIDDEN_TEXT.items():
            if pattern.search(text):
                errors.append(
                    f"PDF page {page_number} contains forbidden {name}"
                )
        width = float(page.mediabox.width)
        height = float(page.mediabox.height)
        if abs(width - 612) > 2 or abs(height - 792) > 2:
            errors.append(
                f"PDF page {page_number} is {width:.1f}x{height:.1f}; expected Letter"
            )
    return errors


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("docx", type=Path)
    parser.add_argument("--pdf", type=Path)
    args = parser.parse_args()
    errors = validate_docx(args.docx)
    if args.pdf:
        errors.extend(validate_pdf(args.pdf))
    if errors:
        print("thesis review document validation: FAIL")
        for error in errors:
            print(f"- {error}")
        return 1
    print("thesis review document validation: PASS")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
