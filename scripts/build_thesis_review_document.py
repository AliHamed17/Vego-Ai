#!/usr/bin/env python3
"""Build a supervisor-review DOCX from the thesis Markdown sources.

The document uses the narrative_proposal preset and includes evidence-gated
front-matter diagrams. It reads the canonical ThesisEvidenceSnapshot-v1 and
never creates or modifies expert labels or VEGO-AI runtime artifacts.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import subprocess
import tempfile
import zipfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parents[1]
SNAPSHOT_PATH = (
    ROOT / "docs/research/thesis-evidence/thesis-evidence-snapshot-v1.json"
)
OUTPUT_DIR = ROOT / "thesis/output"
DEFAULT_PACKAGE_DATE = "2026-07-25"
FIGURE_DIR = ROOT / "thesis/figures/evidence-ready"
FIGURE_MANIFEST = FIGURE_DIR / "figure-assets-v1.json"

CHAPTERS = [
    ROOT / "thesis/chapters/00-abstract.md",
    ROOT / "thesis/chapters/01-introduction.md",
    ROOT / "thesis/chapters/02-background-and-related-work.md",
    ROOT / "thesis/chapters/03-problem-and-research-questions.md",
    ROOT / "thesis/chapters/04-vego-ai-baseline-pipeline.md",
    ROOT / "thesis/chapters/05-human-ai-co-reasoning-artifact.md",
    ROOT / "thesis/chapters/06-evaluation-methodology.md",
    ROOT / "thesis/chapters/07-experimental-results.md",
    ROOT / "thesis/chapters/08-threats-to-validity.md",
    ROOT / "thesis/chapters/09-discussion.md",
    ROOT / "thesis/chapters/10-conclusion-and-phd-continuation.md",
    ROOT / "thesis/chapters/design-theory-governed-reuse.md",
    ROOT / "thesis/chapters/11-references.md",
    ROOT / "thesis/chapters/appendix-a-supplementary.md",
]

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 101, 115)
TABLE_FILL = "F4F6F9"
CALLOUT_FILL = "EEF5FA"
CAUTION_FILL = "FFF6DA"
INK = RGBColor(30, 35, 42)


def git(*args: str) -> str:
    return subprocess.run(
        ["git", "-C", str(ROOT), *args],
        check=True,
        capture_output=True,
        text=True,
    ).stdout.strip()


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def canonical_json_hash(value: object) -> str:
    payload = json.dumps(
        value,
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
    ).encode("utf-8")
    return hashlib.sha256(payload).hexdigest()


def normalize_docx_package(path: Path) -> None:
    """Rewrite the OOXML ZIP with stable ordering and timestamps."""
    normalized_path = path.with_suffix(path.suffix + ".normalized")
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(
        normalized_path,
        "w",
        compression=zipfile.ZIP_STORED,
    ) as target:
        for name in sorted(source.namelist()):
            source_info = source.getinfo(name)
            info = zipfile.ZipInfo(name, date_time=(1980, 1, 1, 0, 0, 0))
            info.compress_type = zipfile.ZIP_STORED
            info.external_attr = source_info.external_attr
            info.create_system = 0
            target.writestr(info, source.read(name))
    normalized_path.replace(path)


def normalize_text(value: str) -> str:
    return (
        value.replace("\u2011", "-")
        .replace("\u2013", "-")
        .replace("\u2014", " - ")
        .replace("\u00a0", " ")
    )


def set_run_font(
    run,
    name: str = "Calibri",
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    tr_pr.append(element)


def prevent_table_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def column_widths(rows: list[list[str]], total: int = 9360) -> list[int]:
    columns = max(len(row) for row in rows)
    scores = []
    for index in range(columns):
        longest = max((len(row[index]) if index < len(row) else 0) for row in rows)
        scores.append(max(8, min(longest, 48)))
    score_total = sum(scores)
    widths = [max(900, round(total * score / score_total)) for score in scores]
    difference = total - sum(widths)
    widths[-1] += difference
    if widths[-1] < 900:
        shortage = 900 - widths[-1]
        widths[-1] = 900
        donor = max(range(len(widths) - 1), key=lambda i: widths[i])
        widths[donor] -= shortage
    return widths


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_page_field(paragraph, field: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def setup_document(
    doc: Document,
    revision: str,
    generated_at: str,
) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    generated = datetime.fromisoformat(generated_at.replace("Z", "+00:00"))
    properties = doc.core_properties
    properties.title = "VEGO-AI MSc Thesis - Evidence-Ready Review Draft"
    properties.subject = "Reusable, traceable, leakage-aware human judgment"
    properties.author = "VEGO-AI Research Project"
    properties.last_modified_by = "VEGO-AI Research Project"
    properties.created = generated
    properties.modified = generated
    properties.comments = (
        "Generated deterministically from canonical revision " + revision
    )

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.widow_control = True

    heading_specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    if "Code Block" not in styles:
        code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = styles["Code Block"]
    code_style.font.name = "Consolas"
    code_style._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code_style.font.size = Pt(8.5)
    code_style.paragraph_format.left_indent = Inches(0.22)
    code_style.paragraph_format.right_indent = Inches(0.12)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(6)
    code_style.paragraph_format.line_spacing = 1.0

    if "Evidence Callout" not in styles:
        callout_style = styles.add_style("Evidence Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout_style = styles["Evidence Callout"]
    callout_style.font.name = "Calibri"
    callout_style.font.size = Pt(10.5)
    callout_style.font.color.rgb = NAVY
    callout_style.paragraph_format.left_indent = Inches(0.18)
    callout_style.paragraph_format.right_indent = Inches(0.12)
    callout_style.paragraph_format.space_before = Pt(6)
    callout_style.paragraph_format.space_after = Pt(8)
    callout_style.paragraph_format.line_spacing = 1.2

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("VEGO-AI MSc Thesis Review Draft")
    set_run_font(run, size=8.5, color=MUTED, bold=True)
    run = p.add_run(f"   |   source {revision[:8]}")
    set_run_font(run, size=8.5, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Evidence-ready draft   |   Page ")
    set_run_font(run, size=8.5, color=MUTED)
    add_page_field(p, "PAGE")
    run = p.add_run(" of ")
    set_run_font(run, size=8.5, color=MUTED)
    add_page_field(p, "NUMPAGES")

    props = doc.core_properties
    props.title = "Reusable Human Judgment in AI-Assisted Domain Model Assessment"
    props.subject = "Evidence-ready VEGO-AI MSc thesis review draft"
    props.author = "VEGO-AI research project"
    props.keywords = "VEGO-AI, reusable human judgment, evidence gate, MSc thesis"


def paragraph_fill(paragraph, fill: str, border: str | None = None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    if border:
        p_bdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), border)
        p_bdr.append(left)
        p_pr.append(p_bdr)


INLINE_RE = re.compile(
    r"(\*\*.+?\*\*|`.+?`|\*[^*].+?\*|\[[^\]]+\]\([^)]+\))"
)


def add_inline(paragraph, text: str) -> None:
    text = normalize_text(text)
    cursor = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=9)
            run.font.color.rgb = DARK_BLUE
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("["):
            label, target = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            run = paragraph.add_run(f"{label} ({target})")
            run.font.color.rgb = BLUE
            run.font.underline = True
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def add_table(doc: Document, rows: list[list[str]]) -> None:
    columns = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    widths = column_widths(rows)
    set_table_geometry(table, widths)
    for r_index, row_data in enumerate(rows):
        row = table.rows[r_index]
        prevent_table_row_split(row)
        if r_index == 0:
            set_repeat_table_header(row)
        for c_index in range(columns):
            cell = row.cells[c_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r_index == 0:
                shade_cell(cell, TABLE_FILL)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
                if len(row_data[c_index]) < 18
                else WD_ALIGN_PARAGRAPH.LEFT
            ) if c_index < len(row_data) else WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            # Keep compact research tables together when they fit on the
            # current or next page. Rows are already marked cantSplit; this
            # additional chain prevents a header or single trailing row from
            # being stranded across a page boundary.
            paragraph.paragraph_format.keep_with_next = r_index < len(rows) - 1
            value = row_data[c_index] if c_index < len(row_data) else ""
            add_inline(paragraph, value)
            for run in paragraph.runs:
                set_run_font(
                    run,
                    size=8.3 if columns >= 5 else 9,
                    color=INK,
                    bold=r_index == 0,
                )
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def is_separator(line: str) -> bool:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def parse_table(lines: list[str], index: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    while index < len(lines) and lines[index].lstrip().startswith("|"):
        if not is_separator(lines[index]):
            rows.append(
                [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
            )
        index += 1
    return rows, index


def add_markdown(
    doc: Document,
    path: Path,
    first_source: bool = False,
    page_break_before: bool = False,
) -> None:
    lines = path.read_text(encoding="utf-8").splitlines()
    paragraph_buffer: list[str] = []
    in_code = False
    code_lines: list[str] = []
    first_level_one_heading = True

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if paragraph_buffer:
            paragraph = doc.add_paragraph()
            add_inline(paragraph, " ".join(part.strip() for part in paragraph_buffer))
            paragraph_buffer = []

    index = 0
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if stripped.startswith("```"):
            flush_paragraph()
            if in_code:
                paragraph = doc.add_paragraph(style="Code Block")
                paragraph_fill(paragraph, "F4F6F9")
                run = paragraph.add_run(normalize_text("\n".join(code_lines)))
                set_run_font(run, name="Consolas", size=8.5, color=INK)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if not stripped:
            flush_paragraph()
            index += 1
            continue
        if stripped.startswith("|") and index + 1 < len(lines) and is_separator(lines[index + 1]):
            flush_paragraph()
            rows, index = parse_table(lines, index)
            add_table(doc, rows)
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            level = len(heading.group(1))
            paragraph = doc.add_paragraph(style=f"Heading {level}")
            if level == 1 and first_level_one_heading:
                paragraph.paragraph_format.page_break_before = page_break_before
                first_level_one_heading = False
            add_inline(paragraph, heading.group(2))
            index += 1
            continue
        if stripped.startswith("> ```"):
            # Diagram source is tracked in thesis/figures and is already named
            # by the preceding figure callout. Keep the review document visual
            # and omit the quoted Mermaid source block itself.
            flush_paragraph()
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("> ```"):
                index += 1
            if index < len(lines):
                index += 1
            continue
        if stripped == ">":
            flush_paragraph()
            index += 1
            continue
        if stripped.startswith("> "):
            flush_paragraph()
            quote_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith("> "):
                quote_lines.append(lines[index].strip()[2:].strip())
                index += 1
            paragraph = doc.add_paragraph(style="Evidence Callout")
            paragraph_fill(paragraph, CALLOUT_FILL, "2E74B5")
            add_inline(paragraph, " ".join(quote_lines))
            continue
        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet:
            flush_paragraph()
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline(paragraph, bullet.group(1))
            index += 1
            continue
        numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if numbered:
            flush_paragraph()
            paragraph = doc.add_paragraph(style="List Paragraph")
            paragraph.paragraph_format.left_indent = Inches(0.28)
            paragraph.paragraph_format.first_line_indent = Inches(-0.18)
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.add_run(f"{numbered.group(1)}. ")
            add_inline(paragraph, numbered.group(2))
            index += 1
            continue
        if re.fullmatch(r"[-*_]{3,}", stripped):
            flush_paragraph()
            index += 1
            continue
        paragraph_buffer.append(stripped)
        index += 1
    flush_paragraph()


FIGURE_FILES = {
    "ladder": "b0-b5-evidence-ladder.png",
    "architecture": "human-judgment-architecture.png",
    "evidence": "current-evidence-profile.png",
    "roadmap": "exp019-027-roadmap.png",
}


def figure_payload(data: dict) -> dict:
    """Return only the canonical data rendered inside tracked figure assets."""
    evidence_keys = [
        "agent4Patterns",
        "substantialPatterns",
        "occasionalPatterns",
        "undeterminedPatterns",
        "reviewItems",
        "reusableJudgments",
        "memoryAdviceItems",
        "comparisonRows",
        "memoryInformedChanges",
    ]
    return {
        "baselines": [
            {
                key: item[key]
                for key in ("id", "name", "status", "evaluationGate", "data")
            }
            for item in data["baselines"]
        ],
        "experiments": [
            {key: item[key] for key in ("id", "title", "status")}
            for item in data["experiments"]
        ],
        "evidence": {
            key: {
                field: data["evidence"][key][field]
                for field in ("value", "unit", "evidenceClass", "claimBoundary")
            }
            for key in evidence_keys
        },
        "labelGate": {
            key: data["labelGate"][key]
            for key in (
                "candidateRows",
                "suppliedLabels",
                "generalizationSafeLabels",
                "quantitativeMinimum",
                "status",
                "accuracyStatus",
            )
        },
        "runtimeHardening": data["runtimeHardening"],
    }


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    # Figure refreshes are intentional review actions. The reviewed PNG bytes
    # are tracked and hash-verified so DOCX builds do not re-rasterize fonts.
    del bold
    return ImageFont.load_default(size=size)


def wrapped(draw: ImageDraw.ImageDraw, text: str, box_width: int, face) -> list[str]:
    words = normalize_text(text).split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if not current or draw.textlength(candidate, font=face) <= box_width:
            current = candidate
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_text_block(draw, xy, text, width, face, fill, spacing=6):
    x, y = xy
    for line in wrapped(draw, text, width, face):
        draw.text((x, y), line, font=face, fill=fill)
        y += face.size + spacing
    return y


def render_figure_assets(data: dict) -> dict[str, Path]:
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    evidence = data["evidence"]
    gate = data["labelGate"]

    def value(key: str) -> int | float | str:
        return evidence[key]["value"]

    bg = "#07111f"
    panel = "#11243a"
    ink = "#edf4ff"
    muted = "#a9bad0"
    line = "#2b4260"
    colors = ["#68d391", "#68d391", "#f6c85f", "#ff7b7b", "#ff7b7b", "#ff7b7b"]
    paths: dict[str, Path] = {}

    image = Image.new("RGB", (1800, 640), bg)
    draw = ImageDraw.Draw(image)
    draw.text((70, 45), "B0-B5 evidence-maturity ladder", font=font(42, True), fill=ink)
    draw.text((70, 100), "Progress means stronger evidence independence, not guaranteed accuracy.", font=font(23), fill=muted)
    box_w, box_h, gap, y = 258, 335, 24, 190
    for index, item in enumerate(data["baselines"]):
        x = 45 + index * (box_w + gap)
        draw.rounded_rectangle((x, y, x + box_w, y + box_h), 18, fill=panel, outline=colors[index], width=5)
        draw.text((x + 18, y + 18), item["id"], font=font(25, True), fill=colors[index])
        end = draw_text_block(draw, (x + 18, y + 58), item["name"], box_w - 36, font(22, True), ink, 4)
        end = draw_text_block(draw, (x + 18, end + 14), item["status"], box_w - 36, font(18), muted, 3)
        draw_text_block(draw, (x + 18, end + 58), item["evaluationGate"], box_w - 36, font(16), muted, 3)
        if index < 5:
            draw.text((x + box_w + 3, y + box_h / 2 - 16), ">", font=font(34, True), fill="#56d6d1")
    path = FIGURE_DIR / "b0-b5-evidence-ladder.png"
    image.save(path, format="PNG", optimize=False, compress_level=9)
    paths["ladder"] = path

    image = Image.new("RGB", (1800, 900), bg)
    draw = ImageDraw.Draw(image)
    draw.text((70, 45), "Unified runtime journey around a frozen baseline", font=font(42, True), fill=ink)
    runtime = data["runtimeHardening"]
    parity = runtime["parityEvidence"]
    model = runtime["modelBoundary"]
    nodes = [
        (
            "1. Frozen Agent 4 baseline",
            f"{value('agent4Patterns')} classifications; byte and provenance locked",
            "#55a7ff",
        ),
        (
            "2. Legacy M1-M4B-1",
            f"default path; {value('reviewItems')} review items and {value('comparisonRows')} comparison rows",
            "#b59cff",
        ),
        (
            "3. Unified contracts",
            f"version {runtime['contractVersion']}; deterministic adapters preserve public artifacts",
            "#56d6d1",
        ),
        (
            "4. Fail-closed parity",
            f"{parity['artifactCount']} artifacts; {parity['comparisonRowCount']} rows; {parity['classificationChangeCount']} changes",
            "#68d391",
        ),
        (
            "5. Human authority",
            "M-01-M-06 deferred; timeout, conflict, or denial preserves the baseline",
            "#f6c85f",
        ),
        (
            "6. Evidence and model gate",
            f"{gate['generalizationSafeLabels']}/{gate['candidateRows']} safe labels; EXP-029 blocked; {model['defaultModel']} remains default",
            "#ff7b7b",
        ),
    ]
    for index, (title, subtitle, color) in enumerate(nodes):
        row, col = divmod(index, 3)
        x, y = 80 + col * 565, 190 + row * 300
        draw.rounded_rectangle((x, y, x + 460, y + 190), 18, fill=panel, outline=color, width=5)
        draw_text_block(draw, (x + 24, y + 28), title, 410, font(28, True), ink, 4)
        draw_text_block(draw, (x + 24, y + 98), subtitle, 410, font(20), muted, 4)
        if col < 2:
            draw.text((x + 495, y + 70), ">", font=font(48, True), fill="#56d6d1")
        elif row == 0:
            draw.text((x + 205, y + 215), "v", font=font(42, True), fill="#56d6d1")
    draw.rounded_rectangle((170, 790, 1630, 855), 14, outline="#ff7b7b", width=4)
    draw.text((210, 804), "NO RUNTIME, MEMORY, EVALUATION, OR MODEL PATH MAY OVERWRITE B0", font=font(28, True), fill="#ffb5b5")
    path = FIGURE_DIR / "human-judgment-architecture.png"
    image.save(path, format="PNG", optimize=False, compress_level=9)
    paths["architecture"] = path

    image = Image.new("RGB", (1800, 900), bg)
    draw = ImageDraw.Draw(image)
    draw.text((70, 45), "Current evidence profile", font=font(42, True), fill=ink)
    draw.text((70, 100), "Different units are separated; blank performance fields remain blank.", font=font(23), fill=muted)
    charts = [
        (
            "Agent 4 output classes",
            [
                ("Substantial", value("substantialPatterns")),
                ("Occasional", value("occasionalPatterns")),
                ("Undetermined", value("undeterminedPatterns")),
            ],
            value("agent4Patterns"),
            ["#56d6d1", "#f6c85f", "#b59cff"],
        ),
        (
            "Mechanism counts (not additive)",
            [
                ("Review queue", value("reviewItems")),
                ("Memory", value("reusableJudgments")),
                ("Advice", value("memoryAdviceItems")),
                ("Comparison", value("comparisonRows")),
            ],
            value("comparisonRows"),
            ["#f6c85f", "#b59cff", "#56d6d1", "#55a7ff"],
        ),
    ]
    for chart_index, (title, values, maximum, palette) in enumerate(charts):
        x = 70 + chart_index * 865
        draw.rounded_rectangle((x, 175, x + 790, 660), 18, fill=panel, outline=line, width=3)
        draw.text((x + 30, 205), title, font=font(28, True), fill=ink)
        for index, ((label, value), color) in enumerate(
            zip(values, palette, strict=True)
        ):
            y = 285 + index * 82
            draw.text((x + 30, y), label, font=font(20), fill=ink)
            draw.rounded_rectangle((x + 245, y + 2, x + 690, y + 29), 12, fill=bg, outline=line)
            width = 0 if value == 0 else max(8, 440 * value / maximum)
            draw.rounded_rectangle((x + 247, y + 4, x + 247 + width, y + 27), 10, fill=color)
            draw.text((x + 710, y - 3), str(value), font=font(22, True), fill=ink)
    draw.rounded_rectangle((70, 700, 1660, 840), 18, fill=panel, outline="#f6c85f", width=4)
    draw.text(
        (105, 730),
        f"{gate['generalizationSafeLabels']} / {gate['candidateRows']} safe expert labels",
        font=font(38, True),
        fill="#f6c85f",
    )
    draw.text(
        (680, 733),
        f"Accuracy, macro-F1, net correction, and p-value: {gate['accuracyStatus']}",
        font=font(24, True),
        fill=ink,
    )
    draw.text((105, 790), "Next: reviewer calibration -> two blind reviews -> adjudication -> development-only error analysis", font=font(22), fill=muted)
    path = FIGURE_DIR / "current-evidence-profile.png"
    image.save(path, format="PNG", optimize=False, compress_level=9)
    paths["evidence"] = path

    image = Image.new("RGB", (1800, 1050), bg)
    draw = ImageDraw.Draw(image)
    draw.text((70, 45), "EXP-019-EXP-027 gated roadmap", font=font(42, True), fill=ink)
    draw.text((70, 100), "The sequence can stop with a null or harmful result; no stage presumes improvement.", font=font(23), fill=muted)
    for index, item in enumerate(data["experiments"]):
        row, col = divmod(index, 3)
        x, y = 70 + col * 570, 180 + row * 270
        color = "#68d391" if item["status"] == "Evaluation-ready" else "#f6c85f" if "Pending" in item["status"] else "#ff7b7b"
        draw.rounded_rectangle((x, y, x + 490, y + 205), 18, fill=panel, outline=color, width=5)
        draw.text((x + 22, y + 20), item["id"], font=font(23, True), fill=color)
        end = draw_text_block(draw, (x + 22, y + 58), item["title"], 445, font(23, True), ink, 3)
        draw_text_block(draw, (x + 22, end + 12), item["status"], 445, font(17), muted, 3)
        if col < 2:
            draw.text((x + 515, y + 70), ">", font=font(42, True), fill="#56d6d1")
        elif row < 2:
            draw.text((x + 225, y + 215), "v", font=font(36, True), fill="#56d6d1")
    path = FIGURE_DIR / "exp019-027-roadmap.png"
    image.save(path, format="PNG", optimize=False, compress_level=9)
    paths["roadmap"] = path
    return paths


def write_figure_manifest(data: dict, paths: dict[str, Path]) -> Path:
    manifest = {
        "schemaVersion": "ThesisReviewFigureAssets-v1",
        "sourceSnapshot": SNAPSHOT_PATH.relative_to(ROOT).as_posix(),
        "figureDataHash": canonical_json_hash(figure_payload(data)),
        "rendererHash": sha256_file(Path(__file__).resolve()),
        "files": {
            key: {
                "path": path.relative_to(ROOT).as_posix(),
                "sha256": sha256_file(path),
            }
            for key, path in sorted(paths.items())
        },
        "claimBoundary": (
            "These reviewed figures visualize mechanism and gate state only. "
            "They do not establish accuracy, generalization, reduced effort, "
            "or benchmark superiority."
        ),
    }
    FIGURE_MANIFEST.write_text(
        json.dumps(manifest, indent=2, ensure_ascii=False) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    return FIGURE_MANIFEST


def refresh_figure_assets(data: dict) -> Path:
    paths = render_figure_assets(data)
    return write_figure_manifest(data, paths)


def load_figure_assets(data: dict) -> dict[str, Path]:
    if not FIGURE_MANIFEST.is_file():
        raise FileNotFoundError(
            f"missing reviewed figure manifest: {FIGURE_MANIFEST}; "
            "run with --refresh-figures after reviewing the canonical snapshot"
        )
    manifest = json.loads(FIGURE_MANIFEST.read_text(encoding="utf-8"))
    if manifest.get("schemaVersion") != "ThesisReviewFigureAssets-v1":
        raise ValueError("unsupported thesis review figure manifest")
    expected_data_hash = canonical_json_hash(figure_payload(data))
    if manifest.get("figureDataHash") != expected_data_hash:
        raise ValueError(
            "reviewed figures do not match the canonical evidence snapshot; "
            "refresh and review them before rebuilding the DOCX"
        )
    if manifest.get("rendererHash") != sha256_file(Path(__file__).resolve()):
        raise ValueError(
            "the figure renderer changed after the reviewed assets were created; "
            "refresh and review the figures"
        )
    records = manifest.get("files")
    if not isinstance(records, dict) or set(records) != set(FIGURE_FILES):
        raise ValueError("figure manifest must contain exactly the four reviewed assets")
    paths: dict[str, Path] = {}
    for key, expected_name in FIGURE_FILES.items():
        record = records[key]
        path = ROOT / record["path"]
        if path.name != expected_name or not path.is_file():
            raise FileNotFoundError(f"missing reviewed thesis figure: {path}")
        if sha256_file(path) != record.get("sha256"):
            raise ValueError(f"reviewed thesis figure hash mismatch: {path}")
        paths[key] = path
    return paths


def add_picture_page(doc: Document, title: str, path: Path, caption: str) -> None:
    heading = doc.add_paragraph(style="Heading 1")
    add_inline(heading, title)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(6.5))
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.keep_with_next = False
    run = caption_p.add_run(normalize_text(caption))
    set_run_font(run, size=9, color=MUTED, italic=True)


def add_cover(
    doc: Document,
    data: dict,
    revision: str,
    package_date: str,
) -> None:
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(20)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("MSc Thesis - Evidence-Ready Supervisor Review Draft")
    set_run_font(run, size=11, color=BLUE, bold=True)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(14)
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run(
        "Reusable Human Judgment in\nAI-Assisted Domain Model Assessment"
    )
    set_run_font(run, size=28, color=NAVY, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    run = subtitle.add_run("The VEGO-AI Case")
    set_run_font(run, size=17, color=DARK_BLUE)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(22)
    date_value = datetime.strptime(package_date, "%Y-%m-%d")
    display_date = f"{date_value.day} {date_value.strftime('%B %Y')}"
    run = meta.add_run(
        "Prepared for supervisor review: Iris and Arnon\n"
        f"{display_date}\n"
        f"Source revision {revision[:12]}"
    )
    set_run_font(run, size=11, color=MUTED)
    callout = doc.add_paragraph(style="Evidence Callout")
    callout.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_fill(callout, CAUTION_FILL, "BF9000")
    run = callout.add_run(
        f"Current evidence gate: {data['labelGate']['generalizationSafeLabels']} / "
        f"{data['labelGate']['candidateRows']} generalization-safe expert labels. "
        "Accuracy, macro-F1, net correction, and paired significance are not yet computable."
    )
    set_run_font(run, size=10.5, color=RGBColor(95, 72, 0), bold=True)
    doc.add_page_break()


def add_review_front_matter(doc: Document, data: dict, figures: dict[str, Path]) -> None:
    heading = doc.add_paragraph(style="Heading 1")
    add_inline(heading, "Review status and evidence boundary")
    p = doc.add_paragraph(style="Evidence Callout")
    paragraph_fill(p, CAUTION_FILL, "BF9000")
    add_inline(
        p,
        "This draft demonstrates the reusable-human-judgment mechanism and its "
        "governance. It does not claim improved accuracy, generalization, reduced "
        "human effort, benchmark superiority, clinical performance, or automatic correction.",
    )
    rows = [
        ["Item", "Current state", "Interpretation"],
        ["Accepted iteration", f"{data['programSnapshot']['latestAcceptedIteration']} - {data['programSnapshot']['verdict']}", "Reliability-only; selects no default"],
        ["Safe labels", f"{data['labelGate']['generalizationSafeLabels']} / {data['labelGate']['candidateRows']}", "Accuracy not yet computable"],
        ["Agent 4 patterns", str(data["evidence"]["agent4Patterns"]["value"]), "Frozen baseline outputs"],
        ["Memory-informed changes", f"{data['evidence']['memoryInformedChanges']['value']} / {data['evidence']['comparisonRows']['value']}", "Current policy cannot create an accuracy delta"],
        ["Next human action", "Approve protocol; calibrate and label", "Two reviewers plus adjudication"],
        ["Formal claim gate", "External N >= 30; target 48", "All statistical and safety criteria required"],
    ]
    add_table(doc, rows)
    h = doc.add_paragraph(style="Heading 2")
    add_inline(h, "Document map")
    for item in [
        "Front matter: evidence ladder, architecture, current plots, and experiment roadmap.",
        "Chapters 1-5: problem, literature, baseline, and implemented artifact.",
        "Chapter 6: preregistered evaluation and stopping rules.",
        "Chapter 7: current evidence and deliberately blank performance fields.",
        "Chapters 8-10: validity, conditional interpretation, and conclusion.",
        "References, design-theory supplement, and Appendix A.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        add_inline(p, item)
    doc.add_page_break()
    add_picture_page(
        doc,
        "Evidence ladder",
        figures["ladder"],
        "Figure R1. B0-B5 separates implemented mechanism evidence from human-gated, holdout, and external evidence.",
    )
    doc.add_page_break()
    add_picture_page(
        doc,
        "Architecture and safety boundary",
        figures["architecture"],
        "Figure R2. Legacy and unified paths meet at fail-closed parity; human, evidence, and model gates remain explicit.",
    )
    doc.add_page_break()
    add_picture_page(
        doc,
        "Current evidence profile",
        figures["evidence"],
        "Figure R3. Descriptive counts and the zero-label stop state; mixed units are not additive.",
    )
    doc.add_page_break()
    add_picture_page(
        doc,
        "Preregistered experiment sequence",
        figures["roadmap"],
        "Figure R4. EXP-019-EXP-027 advances only when the prior evidence and approval gate passes.",
    )
    doc.add_page_break()


def build(
    output_path: Path,
    source_revision: str | None = None,
    package_date: str = DEFAULT_PACKAGE_DATE,
) -> Path:
    data = json.loads(SNAPSHOT_PATH.read_text(encoding="utf-8"))
    revision = source_revision or data["sourceRevision"]
    revision = git("rev-parse", revision)
    if revision != data["sourceRevision"]:
        raise ValueError(
            "requested source revision differs from ThesisEvidenceSnapshot-v1"
        )
    figures = load_figure_assets(data)
    doc = Document()
    setup_document(doc, revision, data["generatedAt"])
    add_cover(doc, data, revision, package_date)
    add_review_front_matter(doc, data, figures)

    for index, chapter in enumerate(CHAPTERS):
        add_markdown(
            doc,
            chapter,
            first_source=index == 0,
            page_break_before=index > 0,
        )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    normalize_docx_package(output_path)
    return output_path


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output", type=Path)
    parser.add_argument("--package-date", default=DEFAULT_PACKAGE_DATE)
    parser.add_argument("--source-revision")
    parser.add_argument(
        "--refresh-figures",
        action="store_true",
        help=(
            "Regenerate the tracked figure assets and their hash manifest from "
            "the canonical snapshot. Review the resulting images before commit."
        ),
    )
    parser.add_argument(
        "--check",
        action="store_true",
        help="Build to temporary storage and require byte equality with --output.",
    )
    args = parser.parse_args()
    if args.refresh_figures:
        data = json.loads(SNAPSHOT_PATH.read_text(encoding="utf-8"))
        manifest = refresh_figure_assets(data)
        print(manifest)
        return 0
    output = (
        args.output
        or OUTPUT_DIR
        / f"VEGO-AI-MSc-Thesis-Evidence-Ready-Draft-{args.package_date}.docx"
    ).resolve()
    if args.check:
        if not output.is_file():
            print(f"missing thesis review document: {output}")
            return 1
        with tempfile.TemporaryDirectory(prefix="vego-thesis-docx-") as temp_dir:
            candidate = Path(temp_dir) / output.name
            build(candidate, args.source_revision, args.package_date)
            if sha256_file(candidate) != sha256_file(output):
                print(f"STALE: {output}")
                return 1
        print(f"FRESH: {output}")
        return 0
    result = build(output, args.source_revision, args.package_date)
    print(result)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
